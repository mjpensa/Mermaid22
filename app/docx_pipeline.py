from docx import Document
from docx.shared import Inches
import os
import tempfile
import subprocess
from typing import Callable, Optional
from detect import detect_mermaid

def render_mermaid(code: str, format: str, width: float) -> str:
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
        f.write(code)
        input_file = f.name
    output_file = tempfile.mktemp(suffix='.' + format)
    cmd = ['node', 'renderer/render_mermaid.js', input_file, output_file, str(width), format]
    subprocess.run(cmd, check=True)
    os.unlink(input_file)
    return output_file

def process_docx(input_path: str, output_path: str, log_callback: Optional[Callable] = None):
    doc = Document(input_path)
    for para in doc.paragraphs:
        text = para.text
        mermaid_code = detect_mermaid(text)
        if mermaid_code:
            png_path = render_mermaid(mermaid_code, 'png', 500)  # assume width
            p = para._element
            p.getparent().remove(p)
            new_para = doc.add_paragraph()
            run = new_para.add_run()
            run.add_picture(png_path, width=Inches(5))
            os.unlink(png_path)
            if log_callback:
                log_callback(text[:120], mermaid_code)
    temp_docx = tempfile.mktemp(suffix='.docx')
    doc.save(temp_docx)
    cmd = ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(output_path), temp_docx]
    subprocess.run(cmd, check=True)
    pdf_file = temp_docx.replace('.docx', '.pdf')
    os.rename(pdf_file, output_path)
    os.unlink(temp_docx)